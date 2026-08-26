import html as html_mod
import sqlite3

from PySide6.QtCore import QDate
from PySide6.QtWidgets import (
    QCheckBox, QDateEdit, QHBoxLayout, QLabel, QMessageBox, QTextBrowser, QVBoxLayout, QWidget,
)

from qfluentwidgets import CardWidget, PushButton, SubtitleLabel

from app.collector.fetch import load_sources
from app.filter.entities import build_entity_matcher, build_entity_terms, load_customers
from app.filter.keywords import filter_db_rows
from app.gui.settings_store import load_settings
from app.gui.workers import DeepSeekWorker
from app.paths import BASE_DIR

CONFIG = BASE_DIR / "config" / "sources.json"
CUSTOMERS = BASE_DIR / "config" / "customers.json"
DB = BASE_DIR / "data" / "intel.db"
EXPORTS = BASE_DIR / "data" / "exports"


class WorkPage(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.worker = None
        self.customers = load_customers(str(CUSTOMERS))
        self.entity_re = build_entity_matcher(build_entity_terms(self.customers))
        self.source_cats = {s.id: s.category for s in load_sources(str(CONFIG))}
        self._build_ui()
        self.refresh_buttons()

    def _build_ui(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(20, 20, 20, 20)
        lay.setSpacing(12)
        lay.addWidget(SubtitleLabel("工作页面"))

        row = QHBoxLayout()
        row.addWidget(QLabel("时间范围:"))
        self.dt_start = QDateEdit()
        self.dt_start.setCalendarPopup(True)
        self.dt_start.setDate(QDate(2025, 1, 1))
        row.addWidget(self.dt_start)
        row.addWidget(QLabel("至"))
        self.dt_end = QDateEdit()
        self.dt_end.setCalendarPopup(True)
        self.dt_end.setDate(QDate.currentDate())
        row.addWidget(self.dt_end)
        # 修改日期不自动汇总：点下方「数据汇总」按钮才生效
        self.chk_dated = QCheckBox("只看有发布日期")
        self.chk_dated.setChecked(False)
        self.chk_dated.toggled.connect(lambda _: self.show_raw())
        row.addWidget(self.chk_dated)
        row.addStretch(1)
        lay.addLayout(row)

        card = CardWidget()
        card_lay = QVBoxLayout(card)
        card_lay.setContentsMargins(16, 16, 16, 16)
        self.chat = QTextBrowser()
        self.chat.setOpenExternalLinks(True)
        self.chat.document().setDefaultStyleSheet(
            "table{border-collapse:collapse;margin:8px 0;}"
            "td,th{border:1px solid #999999;padding:4px 8px;}"
            "th{font-weight:bold;}"
        )
        card_lay.addWidget(self.chat)
        lay.addWidget(card, 1)

        btns = QHBoxLayout()
        self.chk_filter = QCheckBox("只看相关（关键词实时过滤）")
        self.chk_filter.setChecked(True)
        self.chk_filter.setToolTip("勾选后只显示与【电力无线专网/450MHz/频谱/客户】相关的情报：\n命中客户名称或英文国名，或频谱关键词，或电力+专网/通信关键词")
        self.chk_filter.toggled.connect(lambda _: self.show_raw())
        btns.addWidget(self.chk_filter)
        self.btn_raw = PushButton("数据汇总")
        self.btn_raw.setMinimumHeight(34)
        self.btn_raw.clicked.connect(self.show_raw)
        self.btn_analyze = PushButton("分析信息数据")
        self.btn_analyze.setMinimumHeight(34)
        self.btn_analyze.clicked.connect(self.analyze)
        self.btn_export = PushButton("导出内容")
        self.btn_export.setMinimumHeight(34)
        self.btn_export.clicked.connect(self.export)
        btns.addWidget(self.btn_raw)
        btns.addWidget(self.btn_analyze)
        btns.addWidget(self.btn_export)
        lay.addLayout(btns)

    def refresh_buttons(self):
        s = load_settings()
        self.btn_analyze.setVisible(bool(s.get("ai_enabled")))

    def refresh_entities(self):
        """客户清单变化后重建实体匹配（新客户的缩写/国家立即参与相关过滤）。"""
        self.customers = load_customers(str(CUSTOMERS))
        self.entity_re = build_entity_matcher(build_entity_terms(self.customers))
        self.source_cats = {s.id: s.category for s in load_sources(str(CONFIG))}

    def _bubble(self, title, body_html, _color=None):
        """结构化展示：标题粗体、正文卡片、浅色分隔线；错误用红色提示。"""
        h = ""
        if title:
            h += f"<div style='font-weight:600;margin-top:14px;font-size:14px;color:#1f1f1f;'>{html_mod.escape(title)}</div>"
        if _color == "error":
            body_html = f"<span style='color:#d93025;font-weight:600;'>{body_html}</span>"
        h += f"<div style='margin:6px 0 12px 0;padding:10px 12px;background:#fafafa;"
        h += f"border:1px solid #eeeeee;border-radius:8px;'>{body_html}</div>"
        h += "<hr style='border:none;border-top:1px solid rgba(128,128,128,0.2);'>"
        self.chat.append(h)

    def _load_rows(self):
        try:
            conn = sqlite3.connect(str(DB))
            rows = conn.execute(
                "SELECT id, source_id, source_name, title, url, published, summary, country, fetched_at "
                "FROM items ORDER BY (published = '' OR published IS NULL), published DESC, id DESC"
            ).fetchall()
            conn.close()
            return rows
        except Exception:
            return []

    def _date_range_ok(self) -> bool:
        """开始日期不能晚于结束日期；无效时给出提示并返回 False。"""
        if self.dt_start.date() > self.dt_end.date():
            self._bubble("", "时间范围无效：开始日期晚于结束日期，请重新设置后点击「数据汇总」。", _color="error")
            return False
        return True

    def show_raw(self):
        """数据汇总：直接呈现原文，不分析不删减。"""
        self.chat.clear()
        if not self._date_range_ok():
            return
        rows = self._load_rows()
        if not rows:
            self._bubble("", "数据库暂无数据，请先到「数据源」模块点击「立即抓取」。")
            return
        if self.chk_filter.isChecked():
            rows = filter_db_rows(rows, self.entity_re, self.source_cats)
        start_s = self.dt_start.date().toString("yyyy-MM-dd")
        end_s = self.dt_end.date().toString("yyyy-MM-dd")
        if self.chk_dated.isChecked():
            rows = [r for r in rows if r[5] and start_s <= r[5][:10] <= end_s]
        else:
            rows = [r for r in rows if (not r[5]) or (start_s <= r[5][:10] <= end_s)]
        custom = [k.lower() for k in load_settings().get("custom_keywords", []) if k.strip()]
        if custom:
            rows = [r for r in rows if any(k in (r[3] or "").lower() or k in (r[6] or "").lower() for k in custom)]
        cn_map = {s.name: (s.name_cn or s.name) for s in load_sources(str(CONFIG))}
        mode = "相关" if self.chk_filter.isChecked() else "全部原始"
        self._bubble("", f"共 {len(rows)} 条{mode}数据（原文未删减，关键词实时过滤）：")
        for row in rows[:100]:
            src, title, url, summary = row[2], row[3], row[4], row[6]
            if row[5]:
                time_str = "发布：" + row[5][:10]
            else:
                time_str = ""
            src_cn = html_mod.escape(cn_map.get(src, src))
            t_title = html_mod.escape(title)
            t_url = html_mod.escape(url)
            t_sum = html_mod.escape(summary)
            body = (
                f"<span style='background:#ececec;color:#444444;border-radius:4px;"
                f"padding:2px 8px;font-size:12px;'>{src_cn}</span> "
                f"<span style='color:#888888;font-size:12px;'>{html_mod.escape(time_str)}</span><br/>"
                f"<b style='font-size:14px;color:#1f1f1f;'>{t_title}</b>"
            )
            if url.startswith("http"):
                if "news.google.com/rss/articles" in url:
                    link_text = "🔗 查看原文（Google 跳转）"
                else:
                    link_text = "🔗 查看原文"
                body += (
                    f"<br/><a style='color:#1a73e8;text-decoration:none;font-size:13px;' "
                    f"title='{t_url}' href='{t_url}'>{link_text}</a>"
                )
            if summary:
                body += f"<br/><span style='color:#555555;'>{t_sum}</span>"
            self._bubble("", body)

    def analyze(self):
        """分析信息数据：调用 AI 生成总结分析（防重复触发）。"""
        if self.worker and self.worker.isRunning():
            self._bubble("", "分析正在进行中，请稍候……（分析完成前请勿重复点击）")
            return
        s = load_settings()
        if not s.get("ai_enabled"):
            self._bubble("", "AI 功能未启用，请到「AI 设置」模块开启。")
            return
        if not s.get("api_key"):
            self._bubble("", "尚未配置 API Key，请到「AI 设置」模块填写后重试。")
            return
        if not self._date_range_ok():
            return
        rows = self._load_rows()
        if self.chk_filter.isChecked():
            rows = filter_db_rows(rows, self.entity_re, self.source_cats)
        start_s = self.dt_start.date().toString("yyyy-MM-dd")
        end_s = self.dt_end.date().toString("yyyy-MM-dd")
        if self.chk_dated.isChecked():
            rows = [r for r in rows if r[5] and start_s <= r[5][:10] <= end_s]
        else:
            rows = [r for r in rows if (not r[5]) or (start_s <= r[5][:10] <= end_s)]
        custom = [k.lower() for k in load_settings().get("custom_keywords", []) if k.strip()]
        if custom:
            rows = [r for r in rows if any(k in (r[3] or "").lower() or k in (r[6] or "").lower() for k in custom)]
        if not rows:
            self._bubble("", "当前条件下没有可分析的数据（可调整时间范围、关键词或『只看相关』）。")
            return
        # 先把当前筛选的情报汇总显示出来，让用户清楚分析的对象
        self.show_raw()
        self._bubble("", "以上共 " + str(len(rows)) + " 条为分析对象（最多取前 30 条送入 AI）。正在生成分析，请稍候……")
        cn_map = {s.name: (s.name_cn or s.name) for s in load_sources(str(CONFIG))}
        lines = []
        for row in rows[:30]:
            src_cn = cn_map.get(row[2], row[2])
            lines.append(f"- [{src_cn}] {row[3]} | {row[4]}")
        text = "\n".join(lines)
        messages = [
            {"role": "system", "content": "你是电力行业无线专网情报分析师。基于下面提供的情报条目（每条含来源/标题/链接），产出一份简洁专业的中文情报简报。要求：1) 只基于实际提供的情报，不臆造不扩展；2) 结构为：核心要点 → 逐条分析（引用对应条目） → 趋势研判 → 行动建议；3) 明确标注与电力无线专网、450MHz、频谱、客户进展的直接关联程度；4) 篇幅精炼，可用 Markdown 表格对比。"},
            {"role": "user", "content": text},
        ]
        self._bubble("", "正在生成分析（基于以上汇总数据），请稍候……")
        self.btn_analyze.setEnabled(False)
        self.btn_analyze.setText("分析中…")

        def _finish(*_a):
            self.btn_analyze.setEnabled(True)
            self.btn_analyze.setText("分析信息数据")

        self.worker = DeepSeekWorker(s["api_key"], s.get("model", "deepseek-chat"), messages, self)
        self.worker.done.connect(self._show_analysis)
        self.worker.done.connect(_finish)
        self.worker.failed.connect(lambda e: self._bubble("分析失败", html_mod.escape(str(e))))
        self.worker.failed.connect(_finish)
        self.worker.start()

    def _show_analysis(self, content):
        try:
            import markdown
            html = markdown.markdown(content, extensions=["tables"])
        except Exception:
            html = content.replace("\n", "<br/>")
        self._bubble("分析结果", html)

    def export(self):
        """导出 PDF / Word。"""
        rows = self._load_rows()
        if not rows:
            QMessageBox.information(self, "导出", "数据库暂无数据")
            return
        from datetime import datetime
        EXPORTS.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_path, pdf_path = None, None
        word_err = pdf_err = ""

        try:
            from docx import Document
            doc = Document()
            doc.add_heading("无线专网情报汇总", 0)
            for row in rows:
                src, title, url = row[2], row[3], row[4]
                pub = row[5][:10] if row[5] else ""
                doc.add_heading(title, level=2)
                if pub:
                    doc.add_paragraph(f"来源：{src}  发布：{pub}  链接：{url}")
                else:
                    doc.add_paragraph(f"来源：{src}  链接：{url}")
            word_path = EXPORTS / f"intel_{stamp}.docx"
            doc.save(str(word_path))
        except Exception as e:
            word_err = str(e)

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            from xml.sax.saxutils import escape
            pdf_path = EXPORTS / f"intel_{stamp}.pdf"
            doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
            styles = getSampleStyleSheet()
            story = [Paragraph("无线专网情报汇总", styles["Title"])]
            for row in rows[:100]:
                src, title, url = row[2], row[3], row[4]
                pub = row[5][:10] if row[5] else ""
                story.append(Paragraph(escape(title), styles["Heading2"]))
                if pub:
                    story.append(Paragraph(f"来源：{escape(src)}  发布：{escape(pub)}  链接：{escape(url)}", styles["BodyText"]))
                else:
                    story.append(Paragraph(f"来源：{escape(src)}  链接：{escape(url)}", styles["BodyText"]))
                story.append(Spacer(1, 8))
            doc.build(story)
        except Exception as e:
            pdf_err = str(e)

        msgs = []
        if word_path:
            msgs.append(f"Word：{word_path}")
        else:
            msgs.append("Word 导出失败：" + html_mod.escape(word_err))
        if pdf_path:
            msgs.append(f"PDF：{pdf_path}")
        else:
            msgs.append("PDF 导出失败：" + html_mod.escape(pdf_err))
        self._bubble("导出完成", "<br/>".join(msgs))
