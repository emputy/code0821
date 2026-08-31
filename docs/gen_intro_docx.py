# -*- coding: utf-8 -*-
"""生成《无线专网情报监测系统》介绍与使用说明 Word 版，保存到桌面。"""
import os

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DESKTOP = os.path.join(os.path.expanduser("~"), "OneDrive", "桌面")
OUT = os.path.join(DESKTOP, "无线专网情报监测系统_介绍与使用说明.docx")

doc = Document()

# 全局中文字体
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

NAVY = RGBColor(0x1F, 0x3B, 0x57)
BLUE = RGBColor(0x2C, 0x5F, 0x8A)


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    r = p.add_run(text)
    set_font(r, 15, True, NAVY)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 12, True, BLUE)
    return p


def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


# ---------- 标题 ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("无线专网情报监测系统")
set_font(r, 22, True, NAVY)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("（450MHz / 配电无线专网行业情报）介绍、功能、使用与改进建议")
set_font(r, 11, False, RGBColor(0x66, 0x66, 0x66))

# ---------- 一、程序简介 ----------
h1("一、程序简介")
body("本系统面向 450MHz 无线专网 / 电力无线专网领域（配电与电力行业、友商动态、重点国家频谱监管、客户项目进展），"
     "提供“自动采集 → 智能过滤 → 情报整理 → AI 分析 → 报告导出”的一站式情报监测能力。")
h2("技术特点")
for b in [
    "Python 3.12 + PySide6 + qfluentwidgets：现代化桌面界面（左侧导航 + 卡片式布局），Windows 原生体验。",
    "SQLite 本地存储：data/intel.db 全量入库，界面实时筛选，无需外部数据库。",
    "多类型数据源抓取：RSS 订阅、HTML 页面链接提取、Sitemap 站点地图三种方式，6 线程并发，带超时与反爬应对。",
    "DeepSeek 大模型接入：对情报条目自动生成中文简报（核心要点 / 逐条分析 / 趋势研判 / 行动建议）。",
    "可打包分发：PyInstaller 打包为单文件 exe，无需安装 Python 环境即可在其他电脑运行。",
]:
    bullet(b)

# ---------- 二、功能清单 ----------
h1("二、功能清单")
funcs = [
    ["情报采集", "并发抓取 12 个数据源（450MHz 联盟、诺基亚/爱立信/中兴、8 家客户官网），约 30~60 秒完成一轮；支持按源走代理/直连、SSL 校验开关、URL 包含/排除过滤。"],
    ["数据存储", "SQLite 全量入库；采集日志与错误信息实时回显；异常写入 data/error.log 便于排查。"],
    ["情报工作台", "按关键词、客户阶段（5 阶段）、国家、时间范围筛选；相关度开关（只显示与电力专网/450MHz/频谱/客户相关的条目）；点击“查看原文”直达来源页。"],
    ["客户全景管理", "内置 5 阶段客户清单（频谱洞察→意愿明确→POC 测试→频谱申请→已获取频谱）；可新增客户，自动联动实体匹配与可视化。"],
    ["AI 智能分析", "配置 DeepSeek API Key 后，一键对筛选结果生成结构化中文情报简报（Markdown 渲染）；支持模型切换（flash / pro）。"],
    ["可视化图表", "客户阶段全景（各阶段客户数）、客户地区分布、情报来源/类别分布、情报阶段分布（柱状图 + 饼图）。"],
    ["报告导出", "导出 Word（全部条目）与 PDF（前 100 条），自动保存在 data/exports/ 目录。"],
    ["定时采集", "按设定间隔（1~365 天）与开始时间自动触发采集，界面上显示上次/下次运行时间。"],
    ["设置中心", "API Key 密码框保存、连接测试、AI 开关、模型选择、自定义筛选关键词（每行一个）。"],
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "功能模块"
hdr[1].text = "说明"
for name, desc in funcs:
    row = tbl.add_row().cells
    row[0].text = name
    row[1].text = desc
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, 9.5)

# ---------- 三、如何使用 ----------
h1("三、如何使用")
h2("3.1 运行方式")
body("方式一（推荐分发）：打包 exe")
bullet("单文件版：dist\\IntelMonitor.exe —— 拷到任意 Windows 10/11 64 位电脑双击即用，首次运行会在 exe 旁边自动生成 config/ 与 data/ 文件夹。")
bullet("文件夹版：dist\\IntelMonitor-ondir\\ —— exe 与 _internal、config、data 需保持在同一文件夹内使用。")
body("方式二（源码运行）：")
bullet("安装依赖：pip install -r requirements.txt（需 Python 3.10+）")
bullet("启动：python run.py（或 python -m app.gui）")
body("方式三（PowerShell 快速启动）：")
body("Start-Process \"D:\\code0821\\dist\\IntelMonitor.exe\"　# 或 onedir 版 exe 路径")

h2("3.2 首次使用流程")
for s in [
    "① 启动程序 → 左侧导航进入「AI 设置」页面",
    "② 填入 DeepSeek API Key（sk-...），点「测试连接」验证（可先填临时 Key 测试，不会覆盖已保存 Key）",
    "③ 按需设置：启用 AI、选择模型（flash 快速 / pro 更强）、定时采集间隔与开始时间、自定义筛选关键词 → 点「保存设置」",
    "④ 进入「数据源」页面 → 点「采集」按钮 → 等待 30~60 秒（页面显示进度与日志）",
    "⑤ 进入「工作页面」按需筛选情报（关键词/阶段/国家/时间）→ 勾选条目点「分析信息数据」生成 AI 简报 → 「导出」Word/PDF",
    "⑥ 在「可视化」页面查看客户阶段全景、地区分布等图表",
]:
    bullet(s)

h2("3.3 界面导航")
nav = doc.add_table(rows=1, cols=2)
nav.style = "Light Grid Accent 1"
h = nav.rows[0].cells
h[0].text = "导航项"
h[1].text = "用途"
for name, desc in [
    ["工作页面", "情报列表、筛选、AI 分析、报告导出"],
    ["可视化", "客户阶段全景 / 地区分布 / 来源分布 / 阶段分布图表"],
    ["AI 设置", "API Key、模型、定时采集、自定义关键词、立即采集"],
    ["数据源", "数据源管理、采集进度日志、客户全景（5 阶段）管理、情报列表"],
]:
    row = nav.add_row().cells
    row[0].text = name
    row[1].text = desc
for row in nav.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, 9.5)

# ---------- 四、注意事项 ----------
h1("四、需要注意的事项")
for b in [
    "网络与代理：部分数据源（如境外官网）需要代理访问，代理地址配置在 config/sources.json 顶层的 \"proxy\" 字段（当前 127.0.0.1:7890）。代理/VPN 失效时，走代理的源会报超时，直连源不受影响；可在配置中删除 proxy 字段恢复纯直连。",
    "API Key 安全：DeepSeek API Key 保存在本地 config/settings.json，且该文件已被 .gitignore 排除、不会随 exe 打包分发——新电脑上需重新填写自己的 Key。",
    "数据文件：config/（数据源、客户配置）与 data/（intel.db 数据库、exports 导出目录、error.log 日志）默认在程序旁边生成；备份或迁移时请连同整个文件夹一起拷贝。",
    "打包与系统：单文件 exe 首次启动需要解压、稍慢；exe 未做数字签名，Windows SmartScreen 可能提示“未知发布者”，选择“仍要运行”即可；目标电脑需 Windows 10/11 64 位。",
    "采集行为：默认 6 线程并发，对目标站点压力较小；个别官网源可能返回 403（反爬）或内容更新滞后，属正常现象。",
    "异常排查：界面操作报错会自动记录到 data/error.log，可据此定位问题。",
]:
    bullet(b)

# ---------- 五、可改进的方向 ----------
h1("五、可改进的方向")
h2("5.1 数据源覆盖（当前最值得投入）")
for b in [
    "客户官网补齐：30 家客户中目前 8 家官网可抓，其余 22 家因反爬 403（如 DEWA、PPC）、TLS 异常（如 PLN、MEA）、无新闻栏目（如 PITC、SNEL）、JS 渲染（如 CNEL、Edenor）等原因未接入；可引入 Playwright 无头浏览器抓取 JS 站、自建 RSSHub 中转、或改用各国官方通讯社专题页。",
    "聚合源补充：Google News 类聚合搜索源依赖代理且在国内网络下不稳定，可考虑国内可直连的替代聚合或自建 RSS 服务。",
]:
    bullet(b)
h2("5.2 抓取质量与效率")
for b in [
    "标题质量：部分站点（如 EPM）抓到的链接文本是“阅读全文”占位，可增加“标题回退到 URL slug / 页面 <title>”规则。",
    "增量与去重：当前全量入库，可按 URL 哈希做增量抓取与去重，减少重复条目与无效请求。",
    "失败重试：对偶发 5xx（如 Eskom 的 522）增加自动重试与采集成功率统计。",
    "关键词体系：配置中 extra_keywords 字段目前未参与过滤，可启用为源级关键词加权。",
]:
    bullet(b)
h2("5.3 智能化与交互")
for b in [
    "语义层：引入向量检索 / 语义相似度，自动聚类同类情报、识别客户相关事件。",
    "通知推送：定时采集完成后通过邮件 / 企业微信 / 钉钉推送摘要。",
    "图表增强：时间趋势线、热力图、钻取联动等。",
]:
    bullet(b)
h2("5.4 工程与交付")
for b in [
    "exe 定制：自定义图标、数字签名（消除 SmartScreen 提示）、自动更新检查。",
    "数据目录策略：将用户数据改为 %APPDATA% 独立目录，程序目录保持只读，便于升级替换。",
    "日志与遥测：可选上传匿名错误报告，便于远程排查。",
    "多语言：界面国际化（中/英）。",
]:
    bullet(b)

# ---------- 附：数据源清单 ----------
h1("附：当前数据源清单（12 个）")
src_rows = [
    ["450MHz 联盟", "alliance_450", "RSS", "450 MHz Alliance"],
    ["友商", "nokia_newsroom", "Sitemap", "诺基亚 Newsroom"],
    ["友商", "ericsson_press", "RSS", "爱立信新闻中心"],
    ["友商", "zte_news", "HTML", "中兴新闻中心"],
    ["客户官网", "cust_tnb", "HTML", "TNB（马来西亚国家能源）"],
    ["客户官网", "cust_eskom", "RSS", "Eskom（南非国家电力）"],
    ["客户官网", "cust_edm_moz", "HTML", "EDM（莫桑比克电力）"],
    ["客户官网", "cust_edg", "Sitemap", "EDG（几内亚电力）"],
    ["客户官网", "cust_epm", "HTML", "EPM（哥伦比亚麦德林电力）"],
    ["客户官网", "cust_cienergies", "HTML", "CI-ENERGIES（科特迪瓦电力）"],
    ["客户官网", "cust_eehc", "HTML", "EEHC（埃及电力控股）"],
    ["客户官网", "cust_edf", "HTML", "EDF（法国电力）"],
]
t2 = doc.add_table(rows=1, cols=4)
t2.style = "Light Grid Accent 1"
h = t2.rows[0].cells
for i, name in enumerate(["类别", "ID", "类型", "说明"]):
    h[i].text = name
for row_data in src_rows:
    row = t2.add_row().cells
    for i, v in enumerate(row_data):
        row[i].text = v
for row in t2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, 9.5)

doc.save(OUT)
print("Word 已生成:", OUT)
